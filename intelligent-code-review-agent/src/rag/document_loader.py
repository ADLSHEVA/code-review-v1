"""Document loading for RAG indexing."""

import os
import re
from pathlib import Path

from langchain_core.documents import Document


class DocumentLoader:
    """Load and split documents for vector store indexing."""

    CODE_EXTENSIONS = {
        ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".go", ".rs",
        ".c", ".cpp", ".h", ".hpp", ".rb", ".php", ".swift", ".kt",
    }

    DOC_EXTENSIONS = {".md", ".txt", ".rst", ".adoc", ".pdf", ".docx"}

    @classmethod
    def load_markdown(cls, file_path: str) -> list[Document]:
        """Load a markdown file, split by headers."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        sections = cls._split_by_headers(content)
        documents = []

        for i, (title, body) in enumerate(sections):
            if not body.strip():
                continue
            doc = Document(
                page_content=body.strip(),
                metadata={
                    "source": file_path,
                    "section_title": title,
                    "section_index": i,
                    "doc_type": "guideline",
                    "language": cls._detect_doc_language(file_path, content),
                },
            )
            documents.append(doc)

        return documents

    @classmethod
    def load_code_file(cls, file_path: str) -> list[Document]:
        """Load a code file as a single document."""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        if not content.strip():
            return []

        ext = Path(file_path).suffix
        language = {
            ".py": "python", ".js": "javascript", ".jsx": "javascript",
            ".ts": "typescript", ".tsx": "typescript", ".java": "java",
            ".go": "go", ".rs": "rust", ".c": "c", ".cpp": "cpp",
        }.get(ext, "unknown")

        return [Document(
            page_content=content,
            metadata={
                "source": file_path,
                "doc_type": "code",
                "language": language,
                "filename": os.path.basename(file_path),
            },
        )]

    @classmethod
    def load_directory(
        cls,
        dir_path: str,
        patterns: list[str] | None = None,
        max_files: int = 100,
    ) -> list[Document]:
        """Recursively load documents from a directory."""
        documents = []
        count = 0

        for root, dirs, files in os.walk(dir_path):
            # Skip hidden directories and common non-essential dirs
            dirs[:] = [
                d for d in dirs
                if not d.startswith(".") and d not in {
                    "__pycache__", "node_modules", ".git", "venv", ".venv",
                    "dist", "build", ".eggs",
                }
            ]

            for filename in files:
                if count >= max_files:
                    break

                filepath = os.path.join(root, filename)
                ext = Path(filename).suffix.lower()

                if ext in cls.DOC_EXTENSIONS:
                    docs = cls.load_markdown(filepath)
                    documents.extend(docs)
                    count += 1
                elif ext in cls.CODE_EXTENSIONS:
                    docs = cls.load_code_file(filepath)
                    documents.extend(docs)
                    count += 1

        return documents

    @classmethod
    def load_pdf(cls, file_path: str) -> list[Document]:
        """Load a PDF file, extracting text page by page."""
        try:
            from pypdf import PdfReader
        except ImportError:
            return []

        reader = PdfReader(file_path)
        documents = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if not text or not text.strip():
                continue
            chunks = cls._chunk_text(text, max_chars=1000)
            for i, chunk in enumerate(chunks):
                documents.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": str(Path(file_path).resolve()),
                        "doc_type": "guideline",
                        "language": "general",
                        "page_number": page_num + 1,
                        "chunk_index": i,
                    },
                ))
        return documents

    @classmethod
    def load_docx(cls, file_path: str) -> list[Document]:
        """Load a DOCX file, extracting paragraphs grouped by headings."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            return []

        doc = DocxDocument(file_path)
        documents = []
        current_section = "Document"
        current_text = []

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                if current_text:
                    text = "\n".join(current_text)
                    for i, chunk in enumerate(cls._chunk_text(text, max_chars=1000)):
                        documents.append(Document(
                            page_content=chunk,
                            metadata={
                                "source": str(Path(file_path).resolve()),
                                "doc_type": "guideline",
                                "language": "general",
                                "section_title": current_section,
                                "chunk_index": i,
                            },
                        ))
                    current_text = []
                current_section = para.text.strip() or current_section
            elif para.text.strip():
                current_text.append(para.text)

        if current_text:
            text = "\n".join(current_text)
            for i, chunk in enumerate(cls._chunk_text(text, max_chars=1000)):
                documents.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": str(Path(file_path).resolve()),
                        "doc_type": "guideline",
                        "language": "general",
                        "section_title": current_section,
                        "chunk_index": i,
                    },
                ))

        return documents

    @classmethod
    def load_single_file(cls, file_path: str, doc_type: str = "guideline") -> list[Document]:
        """Load a single file and route to the appropriate loader."""
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            docs = cls.load_pdf(file_path)
        elif ext == ".docx":
            docs = cls.load_docx(file_path)
        elif ext in cls.DOC_EXTENSIONS:
            docs = cls.load_markdown(file_path)
        elif ext in cls.CODE_EXTENSIONS:
            docs = cls.load_code_file(file_path)
        else:
            return []

        for doc in docs:
            doc.metadata["doc_type"] = doc_type
        return docs

    @staticmethod
    def _chunk_text(text: str, max_chars: int = 1000) -> list[str]:
        """Split text into chunks respecting paragraph boundaries."""
        paragraphs = text.split("\n\n")
        chunks = []
        current = ""
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            if len(current) + len(para) + 2 > max_chars and current:
                chunks.append(current)
                current = para
            else:
                current = current + "\n\n" + para if current else para
        if current:
            chunks.append(current)
        return chunks if chunks else [text[:max_chars]]

    @staticmethod
    def _split_by_headers(content: str) -> list[tuple[str, str]]:
        """Split markdown content by headers."""
        pattern = re.compile(r"^(#{1,6})\s+(.+)$", re.MULTILINE)
        matches = list(pattern.finditer(content))

        if not matches:
            return [("Document", content)]

        sections = []
        for i, match in enumerate(matches):
            title = match.group(2).strip()
            start = match.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(content)
            body = content[start:end].strip()
            sections.append((title, body))

        # Add content before first header
        if matches[0].start() > 0:
            preamble = content[:matches[0].start()].strip()
            if preamble:
                sections.insert(0, ("Introduction", preamble))

        return sections

    @staticmethod
    def _detect_doc_language(file_path: str, content: str) -> str:
        """Try to detect the language a document is about."""
        path_lower = file_path.lower()
        if "python" in path_lower or "py" in path_lower:
            return "python"
        if "javascript" in path_lower or "js" in path_lower:
            return "javascript"
        if "typescript" in path_lower or "ts" in path_lower:
            return "typescript"
        if "security" in path_lower:
            return "security"
        if "plc" in path_lower:
            return "plc"
        return "general"
